# src/data_processing/document_extractor_fixed.py

import os
import json
import re
from datetime import datetime
from typing import Dict, List, Tuple
from docx import Document
from docx.shared import Inches
import pandas as pd

class LegalDocumentExtractor:
    """
    کلاس استخراج متن از اسناد حقوقی ورد - نسخه اصلاح شده
    """
    
    def __init__(self, input_dir: str = "data/raw_documents", output_dir: str = "data/processed_chunks"):
        self.input_dir = input_dir
        self.output_dir = output_dir
        self.metadata_dir = "data/metadata"
        
        # ایجاد فولدرهای خروجی در صورت عدم وجود
        os.makedirs(output_dir, exist_ok=True)
        os.makedirs(self.metadata_dir, exist_ok=True)
        
        # پترن‌های regex برای تشخیص ساختار حقوقی
        self.patterns = {
            'article': re.compile(r'ماده\s*[\u06F0-\u06F9\u0660-\u0669\d]+', re.UNICODE),
            'clause': re.compile(r'بند\s*([الف-ی\d]+)', re.UNICODE),
            'note': re.compile(r'تبصره\s*(\d*)', re.UNICODE),
            'chapter': re.compile(r'فصل\s*(\d+|اول|دوم|سوم|چهارم|پنجم)', re.UNICODE),
            'law_title': re.compile(r'قانون\s+(.+?)(?:\n|$)', re.UNICODE),
            'approval_date': re.compile(r'مصوب\s*(\d{1,2}/\d{1,2}/\d{4})', re.UNICODE),
            'authority': re.compile(r'(مجلس شورای اسلامی|هیئت وزیران|شورای عالی)', re.UNICODE)
        }
    
    def extract_text_from_docx(self, file_path: str) -> Tuple[str, Dict]:
        """
        استخراج متن و metadata از فایل ورد - نسخه بهبود یافته
        """
        try:
            doc = Document(file_path)
            
            # استخراج متن از paragraphs با حفظ ساختار
            paragraphs = []
            text_lines = []
            
            print(f"📄 استخراج از {len(doc.paragraphs)} پاراگراف...")
            
            # پردازش paragraphs
            for i, para in enumerate(doc.paragraphs):
                para_text = para.text.strip()
                if para_text:
                    # حفظ اطلاعات استایل
                    para_info = {
                        'text': para_text,
                        'style': para.style.name if para.style else 'Normal',
                        'is_heading': self._is_heading_style(para.style.name if para.style else ''),
                        'line_number': i
                    }
                    paragraphs.append(para_info)
                    
                    # اضافه کردن به لیست متن با line break
                    text_lines.append(para_text)
                    
                    # اضافه کردن خط خالی بعد از عناوین مهم
                    if self._is_important_heading(para_text):
                        text_lines.append("")  # خط خالی
            
            # پردازش جداول
            tables_text = []
            for table in doc.tables:
                table_text = self._extract_table_text(table)
                if table_text:
                    tables_text.append(table_text)
                    text_lines.extend(table_text.split('\n'))
            
            # ترکیب متن کامل با حفظ line break ها
            full_text = '\n'.join(text_lines)
            
            print(f"📝 استخراج شد: {len(text_lines)} خط، {len(full_text.split())} کلمه")
            
            # استخراج metadata اولیه
            metadata = self._extract_metadata(full_text, file_path)
            
            return full_text, {
                'metadata': metadata,
                'paragraphs': paragraphs,
                'tables': tables_text,
                'total_lines': len(text_lines),
                'extraction_date': datetime.now().isoformat()
            }
            
        except Exception as e:
            print(f"❌ خطا در استخراج فایل {file_path}: {str(e)}")
            return "", {}
    
    def _is_important_heading(self, text: str) -> bool:
        """تشخیص عناوین مهم که بعد از آنها باید خط خالی باشد"""
        important_patterns = [
            r'فصل\s*[\d\u06F0-\u06F9]+',
            r'بخش\s*[\d\u06F0-\u06F9]+',
            r'ماده\s*[\d\u06F0-\u06F9]+',
            r'قانون\s+',
            r'آیین‌نامه\s+',
            r'دستورالعمل\s+'
        ]
        
        for pattern in important_patterns:
            if re.search(pattern, text, re.UNICODE):
                return True
        return False
    
    def _extract_table_text(self, table) -> str:
        """استخراج متن از جداول با فرمت بهتر"""
        table_lines = []
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_cells.append(cell_text)
            if row_cells:
                table_lines.append(' | '.join(row_cells))
        return '\n'.join(table_lines)
    
    def _is_heading_style(self, style_name: str) -> bool:
        """تشخیص سبک عنوان"""
        heading_styles = ['Heading', 'Title', 'Subtitle', 'عنوان', 'heading']
        return any(style.lower() in style_name.lower() for style in heading_styles)
    
    def _extract_metadata(self, text: str, file_path: str) -> Dict:
        """
        استخراج metadata از متن سند - نسخه بهبود یافته
        """
        metadata = {
            'file_name': os.path.basename(file_path),
            'file_path': file_path,
            'document_type': self._detect_document_type(text),
            'title': self._extract_title(text),
            'approval_date': self._extract_approval_date(text),
            'authority': self._extract_authority(text),
            'articles_count': len(self.patterns['article'].findall(text)),
            'chapters_count': len(self.patterns['chapter'].findall(text)),
            'notes_count': len(self.patterns['note'].findall(text)),
            'word_count': len(text.split()),
            'char_count': len(text),
            'line_count': len(text.split('\n')),
            'language': 'persian'
        }
        
        print(f"📊 Metadata: {metadata['articles_count']} ماده، {metadata['line_count']} خط")
        
        return metadata
    
    def _detect_document_type(self, text: str) -> str:
        """تشخیص نوع سند"""
        text_start = text[:1000].lower()  # بررسی 1000 کاراکتر اول
        
        if 'قانون' in text_start:
            return 'قانون'
        elif 'آیین نامه' in text_start or 'آیین‌نامه' in text_start:
            return 'آیین‌نامه'
        elif 'دستورالعمل' in text_start:
            return 'دستورالعمل'
        elif 'مصوبه' in text_start:
            return 'مصوبه'
        elif 'بخشنامه' in text_start:
            return 'بخشنامه'
        else:
            return 'سند حقوقی'
    
    def _extract_title(self, text: str) -> str:
        """استخراج عنوان سند"""
        lines = text.split('\n')[:15]  # بررسی 15 خط اول
        
        for line in lines:
            line = line.strip()
            if len(line) > 15 and len(line) < 300:  # طول مناسب برای عنوان
                # حذف کاراکترهای اضافی
                title = re.sub(r'[#*\-=\[\]]+', '', line).strip()
                if title and not title.isdigit() and 'بخش' not in title[:10]:
                    return title
        
        # اگر عنوان پیدا نشد، سعی در استخراج از پترن قانون
        law_match = self.patterns['law_title'].search(text)
        if law_match:
            return law_match.group(1).strip()
        
        return "بدون عنوان"
    
    def _extract_approval_date(self, text: str) -> str:
        """استخراج تاریخ تصویب"""
        date_match = self.patterns['approval_date'].search(text)
        return date_match.group(1) if date_match else ""
    
    def _extract_authority(self, text: str) -> str:
        """استخراج مرجع تصویب‌کننده"""
        authority_match = self.patterns['authority'].search(text)
        return authority_match.group(1) if authority_match else ""
    
    def clean_text(self, text: str) -> str:
        """
        پاک‌سازی متن از کاراکترهای اضافی - نسخه بهبود یافته
        """
        # حفظ line break های مهم
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # پاک‌سازی هر خط
            cleaned_line = line.strip()
            
            if cleaned_line:  # فقط خطوط غیرخالی
                # حذف فاصله‌های اضافی در داخل خط
                cleaned_line = re.sub(r'\s+', ' ', cleaned_line)
                
                # حذف کاراکترهای کنترلی
                cleaned_line = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x84\x86-\x9f]', '', cleaned_line)
                
                # تبدیل ارقام عربی به فارسی
                arabic_to_persian = str.maketrans('٠١٢٣٤٥٦٧٨٩', '۰۱۲۳۴۵۶۷۸۹')
                cleaned_line = cleaned_line.translate(arabic_to_persian)
                
                cleaned_lines.append(cleaned_line)
        
        # ترکیب مجدد با حفظ ساختار
        return '\n'.join(cleaned_lines)
    
    def process_single_document(self, file_path: str) -> Dict:
        """
        پردازش یک سند منفرد - نسخه بهبود یافته
        """
        print(f"🔄 در حال پردازش: {os.path.basename(file_path)}")
        
        # استخراج متن
        raw_text, doc_info = self.extract_text_from_docx(file_path)
        
        if not raw_text:
            print(f"❌ خطا در استخراج متن از {file_path}")
            return {}
        
        # پاک‌سازی متن
        cleaned_text = self.clean_text(raw_text)
        
        # بررسی کیفیت پاک‌سازی
        lines_before = len(raw_text.split('\n'))
        lines_after = len(cleaned_text.split('\n'))
        
        print(f"🧹 پاک‌سازی: {lines_before} → {lines_after} خط")
        
        # ذخیره متن پاک‌سازی شده
        output_file = os.path.join(
            self.output_dir, 
            f"{os.path.splitext(os.path.basename(file_path))[0]}_cleaned.txt"
        )
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(cleaned_text)
        
        # ذخیره metadata
        metadata_file = os.path.join(
            self.metadata_dir,
            f"{os.path.splitext(os.path.basename(file_path))[0]}_metadata.json"
        )
        
        doc_info['cleaned_text_path'] = output_file
        doc_info['metadata']['cleaned_word_count'] = len(cleaned_text.split())
        doc_info['metadata']['cleaned_char_count'] = len(cleaned_text)
        doc_info['metadata']['cleaned_line_count'] = len(cleaned_text.split('\n'))
        
        with open(metadata_file, 'w', encoding='utf-8') as f:
            json.dump(doc_info, f, ensure_ascii=False, indent=2)
        
        print(f"✅ پردازش کامل شد:")
        print(f"   📄 متن: {output_file}")
        print(f"   📊 Metadata: {metadata_file}")
        print(f"   📈 آمار: {doc_info['metadata']['word_count']} کلمه، {doc_info['metadata']['articles_count']} ماده")
        
        return doc_info
    
    def process_all_documents(self) -> List[Dict]:
        """
        پردازش تمام اسناد موجود در فولدر ورودی
        """
        print(f"🚀 شروع پردازش اسناد از: {self.input_dir}")
        
        processed_docs = []
        
        # پیدا کردن تمام فایل‌های ورد
        word_files = []
        for file_name in os.listdir(self.input_dir):
            if file_name.endswith(('.docx', '.doc')):
                word_files.append(os.path.join(self.input_dir, file_name))
        
        if not word_files:
            print("❌ هیچ فایل ورد در فولدر ورودی یافت نشد!")
            return []
        
        print(f"📚 تعداد فایل‌های یافت شده: {len(word_files)}")
        
        # پردازش هر فایل
        for file_path in word_files:
            try:
                doc_info = self.process_single_document(file_path)
                if doc_info:
                    processed_docs.append(doc_info)
            except Exception as e:
                print(f"❌ خطا در پردازش {file_path}: {str(e)}")
        
        # ایجاد خلاصه کلی
        summary = {
            'total_documents': len(processed_docs),
            'total_words': sum(doc['metadata']['word_count'] for doc in processed_docs),
            'total_articles': sum(doc['metadata']['articles_count'] for doc in processed_docs),
            'total_lines': sum(doc['metadata']['line_count'] for doc in processed_docs),
            'document_types': list(set(doc['metadata']['document_type'] for doc in processed_docs)),
            'processing_date': datetime.now().isoformat()
        }
        
        # ذخیره خلاصه
        summary_file = os.path.join(self.metadata_dir, 'processing_summary.json')
        with open(summary_file, 'w', encoding='utf-8') as f:
            json.dump(summary, f, ensure_ascii=False, indent=2)
        
        print(f"\n📊 خلاصه پردازش:")
        print(f"   📄 تعداد اسناد: {summary['total_documents']}")
        print(f"   📝 مجموع کلمات: {summary['total_words']}")
        print(f"   📋 مجموع مواد: {summary['total_articles']}")
        print(f"   📏 مجموع خطوط: {summary['total_lines']}")
        print(f"   🗂️ انواع اسناد: {', '.join(summary['document_types'])}")
        
        return processed_docs


def main():
    """تابع اصلی برای تست"""
    # ایجاد extractor
    extractor = LegalDocumentExtractor()
    
    # پردازش تمام اسناد
    results = extractor.process_all_documents()
    
    if results:
        print(f"\n🎉 پردازش {len(results)} سند با موفقیت انجام شد!")
    else:
        print("❌ هیچ سندی پردازش نشد!")


if __name__ == "__main__":
    main()